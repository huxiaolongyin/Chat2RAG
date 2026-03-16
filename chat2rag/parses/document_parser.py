import asyncio
import os
import re
from abc import ABC, abstractmethod
from contextlib import contextmanager
from typing import TYPE_CHECKING, List, Tuple

import pandas as pd
from docling.backend.pypdfium2_backend import PyPdfiumDocumentBackend
from docling.datamodel.base_models import InputFormat
from docling.document_converter import DocumentConverter, PdfFormatOption
from docling.pipeline.standard_pdf_pipeline import StandardPdfPipeline
from docling_core.types.doc.document import PictureItem, SectionHeaderItem, TextItem
from docx import Document
from docx.document import Document as DocumentObject

from chat2rag.config import CONFIG
from chat2rag.core.enums import DocumentType
from chat2rag.core.logger import get_logger
from chat2rag.schemas.document import DocumentData, SourceLocation

if TYPE_CHECKING:
    from chat2rag.services.contextual_retrieval import ContextualRetrieval

logger = get_logger(__name__)


class DocumentParser(ABC):
    """文档解析器基类"""

    @abstractmethod
    async def parse(self, file_path: str) -> List[DocumentData]:
        """解析文档，返回分块列表"""
        pass


class PDFParser(DocumentParser):
    """PDF解析器"""

    doc_converter = DocumentConverter(
        allowed_formats=[InputFormat.PDF],
        format_options={
            InputFormat.PDF: PdfFormatOption(
                pipeline_cls=StandardPdfPipeline, backend=PyPdfiumDocumentBackend
            )
        },
    )

    def __init__(self, max_chars: int = 600, overlap: int = 100):
        if overlap >= max_chars:
            raise ValueError(
                f"overlap ({overlap}) must be less than max_chars ({max_chars})"
            )
        self._max_chars = max_chars
        self._overlap = overlap

    def _split_by_chars(self, text: str, max_chars: int, overlap: int) -> List[str]:
        text = text.strip()
        if not text:
            return []
        if len(text) <= max_chars:
            return [text]

        chunks: List[str] = []
        start = 0
        text_len = len(text)

        while start < text_len:
            end = min(text_len, start + max_chars)
            chunks.append(text[start:end])

            if end >= text_len:
                break
            start = end - overlap

        return chunks

    def _flush(
        self,
        file_path: str,
        page_num: int,
        result: List[DocumentData],
        buf: list[str],
    ):
        if buf:
            content = "".join(buf).strip()
            if len(content) <= self._max_chars:
                result.append(
                    DocumentData(
                        doc_type=DocumentType.WORD,
                        content=content,
                        source=SourceLocation(
                            file_path=file_path,
                            page_num=page_num,
                        ),
                        answer=None,
                        parent_doc_id=None,
                        chunk_index=None,
                        external_id=None,
                    )
                )
            else:
                chunks = self._split_by_chars(content, self._max_chars, self._overlap)
                for chunk in chunks:
                    result.append(
                        DocumentData(
                            doc_type=DocumentType.WORD,
                            content=chunk,
                            source=SourceLocation(
                                file_path=file_path,
                                page_num=page_num,
                            ),
                            answer=None,
                            parent_doc_id=None,
                            chunk_index=None,
                            external_id=None,
                        )
                    )
            buf.clear()

    async def parse(self, file_path: str) -> List[DocumentData]:
        """解析 PDF 文档，返回分块列表"""
        result: List[DocumentData] = []
        buf: list[str] = []
        page_no = 0

        doc = await asyncio.to_thread(self.doc_converter.convert, file_path)

        for item, *_ in doc.document.iterate_items():
            try:
                page_no = item.prov[0].page_no
            except:
                page_no = 0

            if isinstance(item, PictureItem):
                continue

            if isinstance(item, SectionHeaderItem):
                self._flush(
                    file_path=file_path,
                    page_num=page_no,
                    result=result,
                    buf=buf,
                )
                buf.append(item.text)
                buf.append("\n")
                continue

            if isinstance(item, TextItem):
                buf.append(item.text)
                continue

        self._flush(
            file_path=file_path,
            page_num=page_no,
            result=result,
            buf=buf,
        )
        return result

    async def parse_with_context(
        self,
        file_path: str,
        context_generator: "ContextualRetrieval",
    ) -> List[DocumentData]:
        """
        解析 PDF 并应用 Contextual Retrieval

        Args:
            file_path: 文件路径
            context_generator: ContextualRetrieval 实例

        Returns:
            追加了上下文的分块列表
        """
        chunks = await self.parse(file_path)
        if not chunks:
            return chunks

        raw_content = await self._extract_raw_content(file_path)
        title, summary = await context_generator.extract_document_summary(raw_content)

        chunk_contents = [chunk.content for chunk in chunks]
        contexts = await context_generator.generate_chunk_contexts_batch(
            title=title,
            summary=summary,
            chunks=chunk_contents,
        )

        for i, chunk in enumerate(chunks):
            context = contexts[i] if i < len(contexts) else None
            if context:
                chunk.content = context_generator.apply_context_to_chunk(
                    content=chunk.content,
                    context=context,
                    title=title,
                )

        return chunks

    async def _extract_raw_content(self, file_path: str, max_chars: int = 2000) -> str:
        """提取 PDF 原始文本用于生成摘要"""
        doc = await asyncio.to_thread(self.doc_converter.convert, file_path)
        text_parts = []
        for item, *_ in doc.document.iterate_items():
            if hasattr(item, "text") and item.text:
                text_parts.append(item.text)
        return "\n".join(text_parts)[:max_chars]


class WordParser(DocumentParser):
    """
    Word文档解析器，实现三阶段分块策略：
    1. 按标题层级分块
    2. 按章节编号分块
    3. 按字符长度分块（带重叠）
    """

    _CN_NUM = "一二三四五六七八九十百千万零〇"
    _SECTION_PATTERNS = [
        rf"^(?P<mark>[{_CN_NUM}]+)、\s*(?P<title>.*)$",
        rf"^(?P<mark>[\(（][{_CN_NUM}]+[\)）])\s*(?P<title>.*)$",
    ]
    _SECTION_RES = [re.compile(p) for p in _SECTION_PATTERNS]

    DEFAULT_MAX_CHARS = 600
    DEFAULT_OVERLAP = 100

    def __init__(
        self,
        max_chars: int = DEFAULT_MAX_CHARS,
        overlap: int = DEFAULT_OVERLAP,
        *,
        preserve_hierarchy: bool = True,
    ):
        if overlap >= max_chars:
            raise ValueError(
                f"overlap ({overlap}) must be less than max_chars ({max_chars})"
            )
        self._max_chars = max_chars
        self._overlap = overlap
        self._preserve_hierarchy = preserve_hierarchy

    async def parse(self, file_path: str) -> List[DocumentData]:
        """解析 Word 文档，返回分块列表"""
        self._validate_file(file_path)
        return await asyncio.to_thread(self._parse_sync, file_path)

    async def parse_with_context(
        self,
        file_path: str,
        context_generator: "ContextualRetrieval",
    ) -> List[DocumentData]:
        """
        解析 Word 并应用 Contextual Retrieval

        Args:
            file_path: 文件路径
            context_generator: ContextualRetrieval 实例

        Returns:
            追加了上下文的分块列表
        """
        chunks = await self.parse(file_path)
        if not chunks:
            return chunks

        raw_content = await self._extract_raw_content(file_path)
        title, summary = await context_generator.extract_document_summary(raw_content)

        chunk_contents = [chunk.content for chunk in chunks]
        contexts = await context_generator.generate_chunk_contexts_batch(
            title=title,
            summary=summary,
            chunks=chunk_contents,
        )

        for i, chunk in enumerate(chunks):
            context = contexts[i] if i < len(contexts) else None
            section = chunk.source.section if chunk.source else None
            if context:
                chunk.content = context_generator.apply_context_to_chunk(
                    content=chunk.content,
                    context=context,
                    title=title,
                    section=section,
                )

        return chunks

    async def _extract_raw_content(self, file_path: str, max_chars: int = 2000) -> str:
        """提取 Word 原始文本用于生成摘要"""
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text:
                text_parts.append(para.text)
        return "\n".join(text_parts)[:max_chars]

    def _heading_level(self, style_name: str) -> int | None:
        for prefix in ("Heading ", "标题 "):
            if style_name.startswith(prefix):
                try:
                    return int(style_name[len(prefix) :].strip())
                except ValueError:
                    return None
        return None

    def _validate_file(self, file_path: str) -> None:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        if not file_path.lower().endswith(".docx"):
            raise ValueError(f"Expected .docx file, got: {file_path}")

    def _match_section_mark(self, line: str) -> Tuple[str, str] | None:
        stripped = line.strip()
        for pattern in self._SECTION_RES:
            match = pattern.match(stripped)
            if match:
                mark = match.group("mark")
                title = (match.group("title") or "").strip()
                return mark, title
        return None

    def _split_by_chars(self, text: str, max_chars: int, overlap: int) -> List[str]:
        text = text.strip()
        if not text:
            return []
        if len(text) <= max_chars:
            return [text]

        chunks: List[str] = []
        start = 0
        text_len = len(text)

        while start < text_len:
            end = min(text_len, start + max_chars)
            chunks.append(text[start:end])

            if end >= text_len:
                break
            start = end - overlap

        return chunks

    def _parse_sync(self, file_path: str) -> List[DocumentData]:
        with self._open_document(file_path) as doc:
            paragraphs = self._extract_paragraphs(doc)

        chunks = self._chunk_by_hierarchy(paragraphs, file_path)
        return self._chunk_by_size(chunks, file_path)

    @contextmanager
    def _open_document(self, file_path: str):
        try:
            doc = Document(file_path)
            yield doc
        finally:
            pass

    def _extract_paragraphs(self, doc: DocumentObject) -> List[Tuple[str, int | None]]:
        paragraphs = []
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue
            level = self._heading_level(para.style.name if para.style else "")
            paragraphs.append((text, level))
        return paragraphs

    def _chunk_by_hierarchy(
        self,
        paragraphs: List[Tuple[str, int | None]],
        file_path: str,
    ) -> List[DocumentData]:
        chunks: List[DocumentData] = []
        heading_stack: List[Tuple[int, str]] = []
        content_buffer: List[str] = []
        current_section: str = ""

        for text, level in paragraphs:
            section_match = self._match_section_mark(text)

            if level is not None:
                if content_buffer:
                    chunks.append(
                        self._make_chunk(
                            content_buffer,
                            file_path,
                            section=current_section or None,
                        )
                    )
                    content_buffer.clear()

                while heading_stack and heading_stack[-1][0] >= level:
                    heading_stack.pop()
                heading_stack.append((level, text))

                current_section = " > ".join(h[1] for h in heading_stack)
            elif section_match:
                if content_buffer:
                    chunks.append(
                        self._make_chunk(content_buffer, file_path, section=None)
                    )
                    content_buffer.clear()

                content_buffer.append(text)
            else:
                if not content_buffer and current_section:
                    content_buffer.append(f"[{current_section}]")
                content_buffer.append(text)

        if content_buffer:
            chunks.append(
                self._make_chunk(
                    content_buffer,
                    file_path,
                    section=current_section or None,
                )
            )

        return chunks

    def _make_chunk(
        self,
        lines: List[str],
        file_path: str,
        section: str | None = None,
    ) -> DocumentData:
        content = "\n".join(lines).strip()
        return DocumentData(
            doc_type=DocumentType.WORD,
            content=content,
            source=SourceLocation(file_path=file_path, section=section),
            answer=None,
            parent_doc_id=None,
            chunk_index=None,
            external_id=None,
        )

    def _chunk_by_size(
        self, chunks: List[DocumentData], file_path: str
    ) -> List[DocumentData]:
        final_chunks: List[DocumentData] = []

        for chunk in chunks:
            if len(chunk.content) <= self._max_chars:
                final_chunks.append(chunk)
            else:
                parts = self._split_by_chars(
                    chunk.content, self._max_chars, self._overlap
                )
                for part in parts:
                    final_chunks.append(
                        DocumentData(
                            doc_type=DocumentType.WORD,
                            content=part,
                            source=SourceLocation(
                                file_path=file_path, section=chunk.source.section
                            ),
                            answer=None,
                            parent_doc_id=None,
                            chunk_index=None,
                            external_id=None,
                        )
                    )

        return final_chunks


class MarkdownParser(DocumentParser):
    """Markdown解析器"""

    async def parse(self, file_path: str) -> List[DocumentData]:
        return []


class QAPairParser(DocumentParser):
    """问答对解析器"""

    async def parse(self, file_path: str) -> List[DocumentData]:
        documents = []
        try:
            if file_path.endswith(".csv"):
                try:
                    df = pd.read_csv(file_path, sep=";")
                except:
                    df = pd.read_csv(file_path, sep=",")
            elif file_path.endswith(".xlsx"):
                df = pd.read_excel(file_path)
            else:
                raise ValueError(f"不支持的文件格式: {file_path}")

        except Exception as e:
            raise ValueError(f"读取文件失败: {str(e)}")

        required_columns = ["问题", "答案"]
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            raise ValueError(f"缺少必要的列: {missing_columns}")

        for idx, row in df.iterrows():
            try:
                question_content = row["问题"]
                answer_content = row["答案"]
                chunks = [
                    DocumentData(
                        doc_type=DocumentType.QUESTION,
                        source=SourceLocation(file_path=file_path),
                        content=question_content,
                        answer=answer_content,
                        parent_doc_id=None,
                        chunk_index=None,
                        external_id=None,
                    ),
                    DocumentData(
                        doc_type=DocumentType.QA_PAIR,
                        source=SourceLocation(file_path=file_path),
                        content=f"{question_content}:{answer_content}",
                        answer=None,
                        parent_doc_id=None,
                        chunk_index=None,
                        external_id=None,
                    ),
                ]
                documents.extend(chunks)

            except Exception as e:
                logger.exception(f"Failed to process row {idx}")
                continue

        return documents


class TSVParser(DocumentParser):
    """TSV解析器 - 格式: id\tcontent"""

    async def parse(self, file_path: str) -> List[DocumentData]:
        documents = []
        try:
            df = pd.read_csv(file_path, sep="\t", header=None, names=["id", "content"])
        except Exception as e:
            raise ValueError(f"读取文件失败: {str(e)}")

        for idx, row in df.iterrows():
            content_val = row["content"]
            if pd.isna(content_val) or str(content_val).strip() == "":
                continue
            documents.append(
                DocumentData(
                    doc_type=DocumentType.TSV,
                    source=SourceLocation(file_path=file_path),
                    content=str(content_val),
                    answer=None,
                    parent_doc_id=None,
                    chunk_index=None,
                    external_id=str(row["id"]),
                )
            )
        return documents
